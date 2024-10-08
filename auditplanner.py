def process_text_file(file_path):
    with open(file_path, 'r') as file:
        lines = file.readlines()

    result = {}
    key = None
    values = []

    for line in lines:
        line = line.strip()  # Remove any leading/trailing whitespace

        # If the line is empty, it's a separator for new sections
        if not line:
            if key:
                result[key] = values  # Assign values to the previous key
            key = None  # Reset key for new section
            values = []  # Reset values list for new section
        elif key is None:
            # The first line of a block is treated as the key
            key = line
        else:
            # Subsequent lines are added as values to the current key
            values.append(line)

    # Add the final key and values if there's no trailing empty line
    if key:
        result[key] = values

    return result


# Example usage
file_path = 'your_text_file.txt'
data = process_text_file(file_path)



# from docx import Document

# def create_docx_from_dict(data, file_name='output.docx'):
#     # Create a new Document
#     doc = Document()

#     # Create a table with 2 columns: one for the key (section) and one for the controls (values)
#     table = doc.add_table(rows=1, cols=1)
#     table.style = 'Table Grid'

#     # Iterate over the dictionary and add rows for each key-value pair
#     for section, controls in data.items():
#         row_cells = table.add_row().cells
        
#         # Add the section (key) in the first column
#         row_cells[0].text = section + '\n'
#         row_cells[0].paragraphs[0].runs[0].bold = True
        
#         # Add all controls (values) in the second column, each on a new line
#         row_cells[0].text += '\n'.join(controls)

#     # Save the document
#     doc.save(file_name)



# create_docx_from_dict(data, 'controls_in_separate_rows.docx')



from docx import Document
from datetime import datetime, timedelta

def extract_info_from_section(section):
    # Split the section string by commas
    parts = section.split(',')
    
    # Extract the title, person in charge, and time taken
    title = parts[0].strip()
    person_in_charge = parts[1].strip()
    time_taken = int(parts[2].strip())  # Assuming the time is in minutes

    return title, person_in_charge, time_taken

def add_time(start_time, minutes):
    # Adds minutes to the starting time and returns the new time
    return start_time + timedelta(minutes=minutes)

def create_docx_from_dict_with_time(data, start_time_str, file_name='output_with_time.docx'):
    # Convert the start time string to a datetime object
    start_time = datetime.strptime(start_time_str, '%H:%M')

    # Create a new Document
    doc = Document()

    # Create a table with 4 columns: section title, person in charge, controls, and time
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Add the header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Section Title'
    header_cells[1].text = 'Person In Charge'
    header_cells[2].text = 'Controls'
    header_cells[3].text = 'End Time'

    # Iterate over the data
    for section, controls in data.items():
        # Extract the title, person in charge, and time taken from the section key
        title, person_in_charge, time_taken = extract_info_from_section(section)

        # Calculate the end time for this section
        end_time = add_time(start_time, time_taken)

        # Add a new row to the table
        row_cells = table.add_row().cells

        # Populate the row with the extracted and calculated information
        row_cells[0].text = title
        row_cells[1].text = person_in_charge
        row_cells[2].text = '\n'.join(controls)  # Add controls, each on a new line
        row_cells[3].text = end_time.strftime('%H:%M')

        # Update the start time for the next section
        start_time = end_time

    # Save the document
    doc.save(file_name)


# Start time provided by the user (24-hour format)
start_time = '09:00'

create_docx_from_dict_with_time(data, start_time, 'controls_with_time.docx')
